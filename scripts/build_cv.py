"""Build the bilingual, one-page ATS resumes from one reproducible source."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "cv-source"
BLUE = RGBColor(41, 69, 255)
INK = RGBColor(23, 23, 20)
MUTED = RGBColor(80, 78, 72)


COPY = {
    # Mantener los datos personales sincronizados con content/site.mjs.
    "es": {
        "role": "Desarrollador Full-stack especializado en IA local, automatización y datos",
        "profile_h": "PERFIL",
        "profile": "Estudiante de Ingeniería de Software que construye productos verificables: aplicaciones full-stack, evaluación de modelos locales, automatización y analítica. Priorizo privacidad, pruebas reproducibles y límites de producto explícitos.",
        "skills_h": "CAPACIDADES TÉCNICAS",
        "skills": [
            "Desarrollo e IA: Next.js, React, TypeScript, FastAPI, Python 3.12, APIs REST, LM Studio, RAG",
            "Datos y automatización: Power BI, DAX, Power Query, Excel, SQL, Prisma, SQLite, n8n, OCR local",
            "Calidad: pytest, Vitest, Playwright, Ruff, mypy, ESLint, CI, documentación y evidencia reproducible",
        ],
        "projects_h": "PROYECTOS SELECCIONADOS",
        "projects": [
            {
                "title": "LocalForge AI Lab",
                "tech": "Next.js · TypeScript · FastAPI · Python · LM Studio · RAG",
                "url": "https://mcasos.github.io/cv-portafolio/projects/localforge-ai-lab/",
                "repo_url": "https://github.com/mcasos/localforge-ai-lab",
                "bullets": [
                    "Construí un laboratorio local-first que compara modelos, normaliza métricas y fallos, recomienda por tarea con criterios auditables y consulta documentos mediante RAG con citas.",
                    "Preparé el candidato local v1.0.0-rc.1 con 131 pytest, 15 Vitest, 26 pruebas de scripts y 6 E2E; el soak de 30.07 min completó 20/20 intentos sin fallos ni bloqueos.",
                ],
            },
            {
                "title": "RutaFactura",
                "tech": "Next.js · TypeScript · Prisma · SQLite · OCR",
                "url": "https://mcasos.github.io/cv-portafolio/projects/rutafactura/",
                "repo_url": "https://github.com/mcasos/rutafactura",
                "bullets": [
                    "Desarrollé un MVP auditable para preparar, revisar y simular facturas de transporte con OCR local, confirmación humana, expediente digital y cobranza.",
                    "Validé cálculos monetarios y flujos con 453 pruebas aprobadas, 89 archivos de prueba y 16 migraciones; SUNAT permanece explícitamente simulado.",
                ],
            },
            {
                "title": "ReporteEnerg",
                "tech": "Power BI · DAX · Power Query · Excel",
                "url": "https://mcasos.github.io/cv-portafolio/projects/reporte-energ/",
                "bullets": [
                    "Diseñé un dashboard ejecutivo de siete páginas que integra demanda, operación, calidad y eventos energéticos en una narrativa orientada a decisiones.",
                    "Entregué modelo Power BI, medidas DAX, tres vistas públicas sanitizadas y un libro Excel complementario descargable.",
                ],
            },
        ],
        "education_h": "EDUCACIÓN",
        "education": "Ingeniería de Software — Universidad Peruana de Ciencias Aplicadas (UPC)  |  2022 — Actualidad",
        "cert_h": "CERTIFICACIONES",
        "cert": "IBM Generative AI Fundamentals (5 cursos) · University of Michigan Python Data Capstone · Google: redes y seguridad, Linux y SQL, automatización de ciberseguridad con Python",
    },
    "en": {
        "role": "Full-stack Developer specializing in local AI, automation, and data",
        "profile_h": "PROFILE",
        "profile": "Software Engineering student building verifiable products across full-stack applications, local-model evaluation, automation, and analytics. I prioritize privacy, reproducible testing, and explicit product boundaries.",
        "skills_h": "TECHNICAL CAPABILITIES",
        "skills": [
            "Development & AI: Next.js, React, TypeScript, FastAPI, Python 3.12, REST APIs, LM Studio, RAG",
            "Data & automation: Power BI, DAX, Power Query, Excel, SQL, Prisma, SQLite, n8n, local OCR",
            "Quality: pytest, Vitest, Playwright, Ruff, mypy, ESLint, CI, documentation, reproducible evidence",
        ],
        "projects_h": "SELECTED PROJECTS",
        "projects": [
            {
                "title": "LocalForge AI Lab",
                "tech": "Next.js · TypeScript · FastAPI · Python · LM Studio · RAG",
                "url": "https://mcasos.github.io/cv-portafolio/en/projects/localforge-ai-lab/",
                "repo_url": "https://github.com/mcasos/localforge-ai-lab",
                "bullets": [
                    "Built a local-first lab that compares models, normalizes metrics and failures, recommends by task with auditable criteria, and queries documents through citation-backed RAG.",
                    "Prepared local candidate v1.0.0-rc.1 with 131 pytest, 15 Vitest, 26 script tests, and 6 E2E checks; a 30.07-minute soak completed 20/20 attempts with no failures or stuck runs.",
                ],
            },
            {
                "title": "RutaFactura",
                "tech": "Next.js · TypeScript · Prisma · SQLite · OCR",
                "url": "https://mcasos.github.io/cv-portafolio/en/projects/rutafactura/",
                "repo_url": "https://github.com/mcasos/rutafactura",
                "bullets": [
                    "Built an auditable MVP to prepare, review, and simulate transport invoices using local OCR, human confirmation, digital records, and collections workflows.",
                    "Validated money calculations and workflows with 453 passing tests, 89 test files, and 16 migrations; SUNAT remains explicitly simulated.",
                ],
            },
            {
                "title": "ReporteEnerg",
                "tech": "Power BI · DAX · Power Query · Excel",
                "url": "https://mcasos.github.io/cv-portafolio/en/projects/reporte-energ/",
                "bullets": [
                    "Designed a seven-page executive dashboard integrating energy demand, operations, quality, and events into a decision-oriented narrative.",
                    "Delivered a Power BI model, DAX measures, three sanitized public views, and a downloadable complementary Excel workbook.",
                ],
            },
        ],
        "education_h": "EDUCATION",
        "education": "Software Engineering — Universidad Peruana de Ciencias Aplicadas (UPC)  |  2022 — Present",
        "cert_h": "CERTIFICATIONS",
        "cert": "IBM Generative AI Fundamentals (5 courses) · University of Michigan Python Data Capstone · Google: network security, Linux and SQL, cybersecurity automation with Python",
    },
}


def font(run, size, bold=False, color=INK, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def hyperlink(paragraph, text, url, size=9.2, bold=False):
    relationship = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2945FF")
    props.append(color)
    if bold:
        props.append(OxmlElement("w:b"))
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(round(size * 2)))
    props.append(size_el)
    run.append(props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    link.append(run)
    paragraph._p.append(link)


def set_spacing(paragraph, before=0, after=0, line=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    set_spacing(paragraph, before=5, after=2)
    run = paragraph.add_run(text)
    font(run, 8.5, bold=True, color=BLUE)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_spacing(paragraph, after=1.5, line=1.0)
    paragraph.paragraph_format.left_indent = Inches(.20)
    paragraph.paragraph_format.first_line_indent = Inches(-.14)
    font(paragraph.add_run(text), 8.6, color=INK)
    return paragraph


def build(lang):
    copy = COPY[lang]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(.46)
    section.bottom_margin = Inches(.42)
    section.left_margin = Inches(.58)
    section.right_margin = Inches(.58)
    section.header_distance = Inches(.2)
    section.footer_distance = Inches(.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    if "Project Title" not in [style.name for style in styles]:
        styles.add_style("Project Title", WD_STYLE_TYPE.PARAGRAPH)

    name = doc.add_paragraph()
    set_spacing(name, after=0)
    font(name.add_run("Miguel André Casós Torre"), 21, bold=True)
    role = doc.add_paragraph()
    set_spacing(role, after=3)
    font(role.add_run(copy["role"]), 10.2, bold=True, color=BLUE)
    contact = doc.add_paragraph()
    set_spacing(contact, after=4)
    font(contact.add_run("Lima, Perú  |  miguel.casos@hotmail.com  |  "), 8.6, color=MUTED)
    hyperlink(contact, "github.com/mcasos", "https://github.com/mcasos", 8.6, True)
    font(contact.add_run("  |  "), 8.6, color=MUTED)
    portfolio = "mcasos.github.io/cv-portafolio"
    hyperlink(contact, portfolio, f"https://{portfolio}", 8.6, True)

    rule = doc.add_paragraph()
    set_spacing(rule, after=3)
    rule_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for key, value in (("val", "single"), ("sz", "16"), ("space", "1"), ("color", "C8FF36")):
        bottom.set(qn(f"w:{key}"), value)
    borders.append(bottom)
    rule_pr.append(borders)

    add_heading(doc, copy["profile_h"])
    profile = doc.add_paragraph()
    set_spacing(profile, after=2, line=1.02)
    font(profile.add_run(copy["profile"]), 8.9)

    add_heading(doc, copy["skills_h"])
    for item in copy["skills"]:
        paragraph = doc.add_paragraph()
        set_spacing(paragraph, after=1)
        label, detail = item.split(":", 1)
        font(paragraph.add_run(label + ":"), 8.45, bold=True)
        font(paragraph.add_run(detail), 8.45)

    add_heading(doc, copy["projects_h"])
    for project in copy["projects"]:
        title = doc.add_paragraph(style="Project Title")
        set_spacing(title, before=2.2, after=.5)
        title.paragraph_format.keep_with_next = True
        hyperlink(title, project["title"], project["url"], 10, True)
        font(title.add_run("  |  " + project["tech"]), 8.15, color=MUTED)
        if project.get("repo_url"):
            font(title.add_run("  |  "), 8.15, color=MUTED)
            hyperlink(title, "GitHub", project["repo_url"], 8.15, True)
        for bullet in project["bullets"]:
            add_bullet(doc, bullet)

    add_heading(doc, copy["education_h"])
    education = doc.add_paragraph()
    set_spacing(education, after=1)
    font(education.add_run(copy["education"]), 8.65, bold=True)

    add_heading(doc, copy["cert_h"])
    cert = doc.add_paragraph()
    set_spacing(cert, after=0, line=1.0)
    font(cert.add_run(copy["cert"]), 8.25)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(footer)
    font(footer.add_run("CV / 2026  ·  " + ("ES" if lang == "es" else "EN")), 7, bold=True, color=MUTED)

    target = OUT / f"Miguel_CV_{lang.upper()}.docx"
    doc.save(target)
    return target


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    for language in ("es", "en"):
        print(build(language))
